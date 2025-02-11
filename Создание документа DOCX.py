from docx import Document
from docx.shared import Inches

guest=['Иван', 'Сергей']

for x in guest:
    document=Document()
    document.add_heading('Приглашение', 0)
    p=document.add_paragraph(x+', спешим сообщить')
    document.add_picture('1.jpg', width=Inches(10.25))
    document.add_page_break()
    document.save(x+'.docx')